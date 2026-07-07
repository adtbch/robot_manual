from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('KRAI 2026 — Button Mapping', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Master Board (ESP32-S3) — Controller PS4 DualShock 4')

# ── Helper ──
def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return table

# ══════════════════════════════════════════════════════════════════════
# 1. JOYSTICK & DPAD
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('1. Joystick & DPAD', level=1)
add_table(doc,
    ['Input', 'Tanpa modifier', '+ R2 hold', '+ Circle hold'],
    [
        ['Analog kiri (LX/LY)', 'Motion vx/vy (ANALOG mode)', 'Armbox motor Y jog (DPAD mode)', 'Armbox FB toggle (DPAD mode)'],
        ['Analog kanan (RX/RY)', 'Yaw ±1°', 'Motor X/Y jog (gripper)', '—'],
        ['DPAD', 'Motion vx/vy (DPAD mode)', 'Armbox motor Y jog (ANALOG mode)', 'Armbox FB toggle (ANALOG mode)'],
    ])

# ══════════════════════════════════════════════════════════════════════
# 2. TOMBOL — MODE MANUAL (mode = 0)
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('2. Tombol — Mode Manual (mode = 0)', level=1)
add_table(doc,
    ['Tombol', 'Edge/Hold', 'Aksi'],
    [
        ['Cross', '—', 'Kosong'],
        ['Circle', 'Hold + arah (edge)', 'Arm box manual:\nKiri→FB toggle R\nKanan→FB toggle L\nAtas→done L\nBawah→done R'],
        ['Square', 'Hold', 'Freeze motion (vx=0, vy=0)'],
        ['Triangle', 'Hold', 'Servo B manual (hanya saat READY_TO_STAB)'],
        ['L1', 'Hold', 'Slow — yaw step 150ms, gripper step 5'],
        ['R1', 'Hold', 'Fast — yaw step 25ms, gripper step 40'],
        ['L2 + Triangle', 'Edge', 'User-defined action (gripper_control)'],
        ['L2 + Circle + analog kiri ke kanan', 'Hold', 'Arm box R: pne ON → ARMBOX_WAIT'],
        ['L2 + Circle + analog kiri ke kiri', 'Hold', 'Arm box L: pne ON → ARMBOX_WAIT'],
        ['L2 + Square', 'Edge', 'User-defined action (armbox_control)'],
        ['L2 + analog kanan', 'Hold', 'Yaw snap kardinal (0°/90°/180°/-90°)'],
        ['L2 + R2', 'Analog ≥250', 'Flash lamp fire'],
        ['R2', 'Hold', 'Gripper mode — motor X/Y (analog kanan) + armbox Y jog (analog kiri/DPAD)'],
        ['OPTIONS', 'Edge (zone 1)', 'SetupZone1 — servo homing + motor Y level 0 + motor X enc 200\n(goto odom hanya jika mode=1)'],
        ['OPTIONS', 'Edge (zone 2)', 'Toggle modeKinematics (goto ↔ kn)'],
    ])

# ══════════════════════════════════════════════════════════════════════
# 3. TOMBOL — MODE OTOMATIS (mode = 1)
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('3. Tombol — Mode Otomatis (mode = 1)', level=1)
add_table(doc,
    ['Tombol', 'Edge/Hold', 'Aksi'],
    [
        ['Cross', '—', 'Kosong'],
        ['Circle', 'Hold + arah (edge)', 'Arm box manual:\nKiri→FB toggle R\nKanan→FB toggle L\nAtas→done L\nBawah→done R'],
        ['Square', 'Hold', 'Freeze motion (vx=0, vy=0)'],
        ['Triangle', 'Hold', 'Servo B manual (hanya saat READY_TO_STAB)'],
        ['L1', 'Hold', 'Slow — yaw step 150ms, gripper step 5'],
        ['R1', 'Hold', 'Fast — yaw step 25ms, gripper step 40'],
        ['L2 + analog kanan', 'Hold', 'Yaw snap kardinal (0°/90°/180°/-90°)'],
        ['L2 + R2', 'Analog ≥250', 'Flash lamp fire'],
        ['R2', 'Hold', 'Gripper mode — motor X/Y (analog kanan) + armbox Y jog (analog kiri/DPAD)'],
        ['OPTIONS', 'Edge (zone 1)', 'SetupZone1 — servo homing + odomGoto(1) + motor Y level 0 + motor X enc 200'],
        ['OPTIONS', 'Edge (zone 2)', 'Toggle modeKinematics (goto ↔ kn)'],
    ])

# ══════════════════════════════════════════════════════════════════════
# 4. INPUT MODE & MISC
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('4. Input Mode & Misc', level=1)
add_table(doc,
    ['Tombol', 'Edge/Hold', 'Aksi'],
    [
        ['SHARE', 'Edge', 'Toggle input mode (DPAD ↔ ANALOG)\n(diblok saat odom record ON)'],
        ['L1+R1+L2+R2', 'Edge', 'Toggle invert input (atas↔bawah, kiri↔kanan)'],
        ['PS', '—', 'Kosong'],
    ])

# ══════════════════════════════════════════════════════════════════════
# 5. ODOM RECORD
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('5. Odom Record', level=1)
add_table(doc,
    ['Tombol', 'Aksi'],
    [
        ['R3 + L3', 'Toggle mode record ON/OFF (press, tanpa hold)'],
    ])

doc.add_heading('Odom Record Combos (mode record ON)', level=2)
add_table(doc,
    ['Combo', 'Target'],
    [
        ['R1 + TOUCHPAD', 'Zone1 slot 0'],
        ['L1 + TOUCHPAD', 'Zone1 slot 1'],
        ['R2 + TOUCHPAD', 'Zone1 slot 2'],
        ['L2 + TOUCHPAD', 'Zone1 slot 3'],
        ['R1 + Triangle', 'Approach slot 0'],
        ['L1 + Square', 'Approach slot 1'],
        ['R2 + Triangle', 'Approach slot 2'],
        ['L2 + Square', 'Approach slot 3'],
        ['Cross + TOUCHPAD', 'Forest WP 2'],
        ['Square + TOUCHPAD', 'Forest WP 6'],
        ['Circle + TOUCHPAD', 'Forest WP 7'],
        ['Triangle + TOUCHPAD', 'Forest WP 11'],
    ])

# ══════════════════════════════════════════════════════════════════════
# 6. CONTROLLER MODE (hardware)
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('6. Controller Mode (hardware)', level=1)
add_table(doc,
    ['Hardware', 'Aksi'],
    [
        ['BOOT button (controller) — double-press 200ms', 'Toggle mode 0 (manual) ↔ 1 (auto)\nLED: putih=auto, coklat=manual\nTersimpan ke NVS'],
        ['BOOT button (master) — double-press 200ms', 'Toggle alliance RED ↔ BLUE\nLED: merah=RED, biru=BLUE\nTersimpan ke NVS'],
    ])

# ══════════════════════════════════════════════════════════════════════
# 7. SPEED & INTERVAL TABLE
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('7. Speed & Interval', level=1)
add_table(doc,
    ['Mode', 'Yaw step', 'Send interval', 'Gripper step'],
    [
        ['Normal', '50ms', '20ms', '15 enc'],
        ['L1 (Slow)', '150ms', '20ms', '5 enc'],
        ['R1 (Fast)', '25ms', '20ms', '40 enc'],
    ])

# ══════════════════════════════════════════════════════════════════════
# 8. MODE TABLE
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('8. Mode Behavior', level=1)
add_table(doc,
    ['Mode', 'R2 OFF', 'R2 ON'],
    [
        ['DPAD', 'DPAD → motion vx/vy\nL-stick → —\nR-stick → yaw', 'DPAD → armbox Y jog\nL-stick → armbox Y jog\nR-stick → motor X/Y'],
        ['ANALOG', 'L-stick → motion vx/vy\nDPAD → —\nR-stick → yaw', 'L-stick → motion vx/vy\nDPAD → armbox Y jog\nR-stick → motor X/Y'],
    ])

# ══════════════════════════════════════════════════════════════════════
# 9. AUTO (TANPA TOMBOL)
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('9. Auto (tanpa tombol)', level=1)
add_table(doc,
    ['Trigger', 'Aksi'],
    [
        ['Proximity R (slave2)', 'Arm box R: pne ON → motor Y level 4 + motor X -255 → grab'],
        ['Proximity L (slave2)', 'Arm box L: pne ON → motor Y level 4 + motor K -255 → grab'],
        ['Motor Y sampai target', 'Hold PWM 50 (anti-gravitasi) + active=false'],
        ['Motor K kena limit', 'Auto stop via motor k 0'],
        ['Proximity gripper (auto)', 'IDLE → CLOSING → UP → STRAIGHTEN → READY_TO_STAB'],
    ])

# ══════════════════════════════════════════════════════════════════════
# 10. MOTOR Y LEVEL
# ══════════════════════════════════════════════════════════════════════
doc.add_heading('10. Motor Y Level (default)', level=1)
add_table(doc,
    ['Level', 'Encoder pulse'],
    [
        ['0', '0'],
        ['1', '300'],
        ['2', '600'],
        ['3', '900'],
        ['4', '1200'],
        ['5', '1500'],
    ])

# Save
output_path = r'C:\Users\NITRO 5\Documents\GitHub\robot_manual\KRAI2026Manual\Button_Mapping_KRAI2026.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
